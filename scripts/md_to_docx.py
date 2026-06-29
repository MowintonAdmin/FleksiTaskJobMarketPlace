"""Convert KNOWLEDGE_TRANSFER.md to a formatted Word document."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH = Path(__file__).parent.parent / "KNOWLEDGE_TRANSFER.md"
OUT_PATH = Path(__file__).parent.parent / "KNOWLEDGE_TRANSFER.docx"


def set_heading_color(paragraph, r, g, b):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_inline_code(paragraph, text):
    run = paragraph.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)


def parse_inline(paragraph, text):
    """Parse inline markdown: **bold**, `code`, plain text."""
    pattern = re.compile(r'(`[^`]+`|\*\*[^*]+\*\*)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            add_inline_code(paragraph, part[1:-1])
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            if part:
                paragraph.add_run(part)


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    lines = MD_PATH.read_text(encoding='utf-8').splitlines()

    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        # Filter separator rows
        data_rows = [r for r in table_rows if not re.match(r'^\|[-| :]+\|$', r.strip())]
        if not data_rows:
            in_table = False
            table_rows = []
            return
        parsed = [[c.strip() for c in row.strip().strip('|').split('|')] for row in data_rows]
        cols = max(len(r) for r in parsed)
        t = doc.add_table(rows=len(parsed), cols=cols)
        t.style = 'Table Grid'
        for i, row in enumerate(parsed):
            for j, cell_text in enumerate(row):
                if j < cols:
                    cell = t.cell(i, j)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    parse_inline(p, cell_text)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    if i == 0:
                        for run in p.runs:
                            run.bold = True
        doc.add_paragraph()
        in_table = False
        table_rows = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                # Flush code block
                if code_lines:
                    p = doc.add_paragraph('\n'.join(code_lines))
                    p.style = 'Normal'
                    p.paragraph_format.left_indent = Inches(0.3)
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(8.5)
                        run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:val'), 'clear')
                    shading.set(qn('w:color'), 'auto')
                    shading.set(qn('w:fill'), 'F3F4F6')
                    p._p.get_or_add_pPr().append(shading)
                    doc.add_paragraph()
                in_code_block = False
                code_lines = []
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table rows
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            flush_table()

        stripped = line.strip()

        # Horizontal rule
        if re.match(r'^---+$', stripped):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', m.group(2))
            text = re.sub(r'`([^`]+)`', r'\1', text)
            hmap = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}
            h = doc.add_heading(text, level=level)
            colors = {1: (0x1A, 0x56, 0xDB), 2: (0x1A, 0x56, 0xDB), 3: (0x37, 0x47, 0x51), 4: (0x37, 0x47, 0x51)}
            set_heading_color(h, *colors.get(level, (0, 0, 0)))
            i += 1
            continue

        # Blockquote
        if stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            parse_inline(p, stripped[2:])
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            i += 1
            continue

        # Bullet list
        if re.match(r'^[-*+]\s+', stripped):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, re.sub(r'^[-*+]\s+', '', stripped))
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s+', stripped):
            p = doc.add_paragraph(style='List Number')
            parse_inline(p, re.sub(r'^\d+\.\s+', '', stripped))
            i += 1
            continue

        # Sub-bullet (indented)
        if re.match(r'^\s{2,}[-*+]\s+', line):
            p = doc.add_paragraph(style='List Bullet 2')
            parse_inline(p, re.sub(r'^\s+[-*+]\s+', '', line))
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        parse_inline(p, stripped)
        i += 1

    if in_table:
        flush_table()

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == '__main__':
    main()
